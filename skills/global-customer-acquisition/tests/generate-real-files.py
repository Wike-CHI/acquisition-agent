"""
Autoresearch 真实文件生成测试
使用 python-docx 和 openpyxl 创建真实的报价单和客户列表
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill
from openpyxl.utils import get_column_letter
import os
from datetime import datetime

# ========== 1. 生成报价单 (.docx) ==========
print("=" * 50)
print("📄 生成报价单 (.docx)")
print("=" * 50)

# 创建文档
doc = Document()

# 标题
title = doc.add_heading('报价单 | Quotation', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 报价单信息
doc.add_paragraph(f'报价单号: QT-{int(datetime.now().timestamp())}')
doc.add_paragraph(f'日期: {datetime.now().strftime("%Y-%m-%d")}')
doc.add_paragraph('有效期: 30 天')

# 客户信息
doc.add_heading('客户信息', level=1)
doc.add_paragraph('客户名称: John Smith')
doc.add_paragraph('公司: Ace Belting Company')
doc.add_paragraph('国家: USA')
doc.add_paragraph('邮箱: john@acebelting.com')

# 产品表格
doc.add_heading('产品清单', level=1)
table = doc.add_table(rows=3, cols=6)
table.style = 'Table Grid'

# 表头
headers = ['序号', '产品名称', '规格型号', '数量', '单价 (USD)', '总价 (USD)']
for i, header in enumerate(headers):
    cell = table.cell(0, i)
    cell.text = header
    cell.paragraphs[0].runs[0].font.bold = True

# 产品数据
products = [
    ['1', '风冷皮带接头机', 'A2FRJ-1200', '1', '$8,500', '$8,500'],
    ['2', '分层机', 'A1-800', '1', '$3,200', '$3,200']
]

for row_idx, product in enumerate(products, 1):
    for col_idx, value in enumerate(product):
        table.cell(row_idx, col_idx).text = value

# 总价
doc.add_paragraph('')
total_para = doc.add_paragraph()
total_para.add_run('总计 (Total): $11,700').bold = True

# 付款条款
doc.add_heading('付款条款', level=1)
doc.add_paragraph('• 付款方式: 30% T/T 定金，70% T/T 发货前付清')
doc.add_paragraph('• 交货期: 15-20 个工作日')
doc.add_paragraph('• 质保期: 1 年')
doc.add_paragraph('• 贸易术语: FOB Shanghai / CIF Destination Port')

# 公司信息
doc.add_heading('供应商信息', level=1)
doc.add_paragraph('温州红龙工业设备制造有限公司 (HOLO)')
doc.add_paragraph('地址: 浙江省瑞安市东山街道望新路188号3幢101室')
doc.add_paragraph('电话: +86 577-66856856')
doc.add_paragraph('手机: {{SALESPERSON_MOBILE}}')
doc.add_paragraph('邮箱: {{SALESPERSON_EMAIL}}')

# 保存
os.makedirs('test-output', exist_ok=True)
docx_path = 'test-output/quotation-ace-belting.docx'
doc.save(docx_path)
print(f'✅ 报价单已生成: {docx_path}')
print('')

# ========== 2. 生成客户列表 (.xlsx) ==========
print("=" * 50)
print("📊 生成客户列表 (.xlsx)")
print("=" * 50)

# 创建工作簿
wb = openpyxl.Workbook()

# ========== Sheet 1: 客户列表 ==========
ws1 = wb.active
ws1.title = "客户列表"

# 表头样式
header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
header_font = Font(color="FFFFFF", bold=True)

# 添加表头
headers = ['公司名', '国家', '联系人', '邮箱', 'ICP评分', '客户等级', '状态', '产品兴趣', '预计采购额', '最后联系', '下次跟进']
for col, header in enumerate(headers, 1):
    cell = ws1.cell(row=1, column=col, value=header)
    cell.fill = header_fill
    cell.font = header_font
    cell.alignment = Alignment(horizontal='center')

# 添加客户数据
customers = [
    ['Ace Belting Company', 'USA', 'John Smith', 'john@acebelting.com', 85, 'A级', '已报价', '风冷机', '$50,000', '2026-03-27', '2026-03-30'],
    ['Sempertrans USA LLC', 'USA', 'Joni Derry', 'joni.derry@semperitgroup.com', 92, 'A级', '谈判中', '水冷机', '$120,000', '2026-03-26', '2026-03-28'],
    ['Dorner Manufacturing', 'USA', 'Mike Johnson', 'mjohnson@dorner.com', 78, 'B级', '有意向', '裁切机', '$35,000', '2026-03-25', '2026-03-29']
]

for customer in customers:
    ws1.append(customer)

# 调整列宽
for col in range(1, len(headers) + 1):
    ws1.column_dimensions[get_column_letter(col)].width = 18

# ========== Sheet 2: Pipeline 概览 ==========
ws2 = wb.create_sheet("Pipeline 概览")

pipeline_data = [
    ['状态', '客户数', '预计金额 (USD)'],
    ['新线索', 0, 0],
    ['已联系', 0, 0],
    ['有意向', 1, 35000],
    ['已报价', 1, 50000],
    ['谈判中', 1, 120000],
    ['已成交', 0, 0],
    ['已流失', 0, 0],
    ['总计', 3, 205000]
]

for row in pipeline_data:
    ws2.append(row)

# 格式化
for col in range(1, 4):
    ws2.cell(row=1, column=col).font = Font(bold=True)
    ws2.cell(row=9, column=col).font = Font(bold=True)

ws2.column_dimensions['A'].width = 12
ws2.column_dimensions['B'].width = 10
ws2.column_dimensions['C'].width = 18

# ========== Sheet 3: 等级分布 ==========
ws3 = wb.create_sheet("等级分布")

grade_data = [
    ['等级', '客户数', '预计金额 (USD)', '占比'],
    ['A级 (≥80)', 2, 170000, '66.7%'],
    ['B级 (60-79)', 1, 35000, '33.3%'],
    ['C级 (40-59)', 0, 0, '0%'],
    ['D级 (<40)', 0, 0, '0%'],
    ['总计', 3, 205000, '100%']
]

for row in grade_data:
    ws3.append(row)

# 格式化
for col in range(1, 5):
    ws3.cell(row=1, column=col).font = Font(bold=True)
    ws3.cell(row=6, column=col).font = Font(bold=True)

for col in ['A', 'B', 'C', 'D']:
    ws3.column_dimensions[col].width = 18

# ========== Sheet 4: 待跟进客户 ==========
ws4 = wb.create_sheet("待跟进")

follow_up_headers = ['公司名', '国家', 'ICP评分', '状态', '最后联系', '下次跟进']
for col, header in enumerate(follow_up_headers, 1):
    cell = ws4.cell(row=1, column=col, value=header)
    cell.font = Font(bold=True)

# 筛选需要跟进的客户（假设今天是 2026-03-28）
follow_up_customers = [
    ['Sempertrans USA LLC', 'USA', 92, '谈判中', '2026-03-26', '2026-03-28']
]

for customer in follow_up_customers:
    ws4.append(customer)

for col in ['A', 'B', 'C', 'D', 'E', 'F']:
    ws4.column_dimensions[col].width = 20

# 保存
xlsx_path = 'test-output/customer-list-2026-03-28.xlsx'
wb.save(xlsx_path)
print(f'✅ 客户列表已生成: {xlsx_path}')
print('')

# ========== 测试汇总 ==========
print("=" * 50)
print("📊 测试汇总")
print("=" * 50)
print(f'✅ 报价单 (.docx): {docx_path}')
print(f'✅ 客户列表 (.xlsx): {xlsx_path}')
print('')
print('文件大小:')
if os.path.exists(docx_path):
    print(f'  📄 报价单: {os.path.getsize(docx_path)} bytes')
if os.path.exists(xlsx_path):
    print(f'  📊 客户列表: {os.path.getsize(xlsx_path)} bytes')
print('')
print('🎉 真实文件生成测试完成！')
