import docx
import os
import re
import json

def extract_parameters_from_docx(file_path, product_name):
    """从Word文档中提取参数"""
    try:
        doc = docx.Document(file_path)
        
        print(f'\n{"=" * 60}')
        print(f'{product_name}')
        print(f'{"=" * 60}')
        print(f'文件: {os.path.basename(file_path)}')
        
        # 提取所有段落文本
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        # 提取表格中的数据
        tables_data = []
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_data.append(cell_text)
                if row_data:
                    table_rows.append(row_data)
            if table_rows:
                tables_data.append(table_rows)
        
        # 打印前30个段落
        print(f'\n文本内容（前30段）:')
        for idx, para in enumerate(paragraphs[:30], 1):
            print(f'{idx}. {para}')
        
        # 打印表格数据
        if tables_data:
            print(f'\n表格数据（{len(tables_data)}个表格）:')
            for table_idx, table in enumerate(tables_data, 1):
                print(f'\n表格 {table_idx}:')
                for row in table[:20]:  # 只打印前20行
                    print(f'  {" | ".join(row)}')
        
        return {
            'paragraphs': paragraphs,
            'tables': tables_data
        }
        
    except Exception as e:
        print(f'错误: {e}')
        return None

# 读取欧式分层机
print('\n' + '=' * 60)
print('📐 欧式分层机参数提取')
print('=' * 60)

plyseparator_files = [
    ('temp-plyseparator-欧式分层机.docx', '欧式分层机'),
    ('temp-plyseparator-A2FCJ130-00-00美式分层机说明书(2).docx', '美式分层机'),
    ('temp-plyseparator-A3FCJ130-00-00立式分层机说明书.docx', '立式分层机'),
    ('temp-plyseparator-750分层机 设计师看过 2025.12.17.docx', '国产750分层机'),
]

for file_name, product_name in plyseparator_files:
    if os.path.exists(file_name):
        extract_parameters_from_docx(file_name, product_name)
